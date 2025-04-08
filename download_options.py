import pandas as pd
import json
from io import BytesIO
from docx import Document
from fpdf import FPDF

def get_json(data):
    return json.dumps(data, indent=4)

def get_csv(data):
    df = pd.DataFrame([data])
    return df.to_csv(index=False)

def get_word(data):
    doc = Document()
    doc.add_heading("Result",level=1)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Result", ln=True, align="C")
    for key, value in data.items():
        pdf.cell(200, 10, txt=f"{key}: {value}", ln=True, align="L")
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

def get_excel(data):
    df = pd.DataFrame([data])
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="sheet1")
    buffer.seek(0)
    return buffer
    